import os
from docx import Document

# 创建一个新的Word文档
doc = Document()

# 定义数据
dataList = []
for i in range(100):
    data = {
        "姓名": i,
        "性别": i,
        "年龄": i
    }
    dataList.append(data)

# 创建"data"目录，如果它不存在
if not os.path.exists("data"):
    os.mkdir("data")

# 向文档中添加内容
for i, data in enumerate(dataList):
    for key, value in data.items():
        # 添加文本
        doc.add_paragraph(f"{key}: {value}")

    # 保存文档到"data"目录下
    filename = os.path.join("data", f"个人信息{i}.docx")
    doc.save(filename)
    # 为下一次迭代清空文档内容
    doc = Document()